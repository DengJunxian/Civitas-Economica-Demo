"""Helpers for generating downloadable policy and replay report artifacts."""

from __future__ import annotations

import json
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def safe_slug(text: str, fallback: str = "report", max_length: int = 36) -> str:
    normalized = re.sub(r"\s+", "-", str(text or "").strip().lower())
    normalized = re.sub(r"[^a-z0-9\-_]+", "-", normalized).strip("-_")
    return (normalized or fallback)[:max_length]


def serializable_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    return json.loads(json.dumps(payload, ensure_ascii=False, default=str))


def official_report_meta(report_type: str, title: str) -> Dict[str, str]:
    now = datetime.now()
    stamp = now.strftime("%Y%m%d-%H%M%S")
    prefix = {
        "policy_lab": "CIVITAS-POL",
        "history_replay": "CIVITAS-HIS",
    }.get(report_type, "CIVITAS-RPT")
    recipient = {
        "policy_lab": "政策研究会商组、相关业务部门",
        "history_replay": "政策评估组、历史复盘与校准团队",
    }.get(report_type, "内部研究与评估团队")
    return {
        "report_no": f"{prefix}-{stamp}",
        "date_cn": now.strftime("%Y年%m月%d日"),
        "generated_at": now.isoformat(timespec="seconds"),
        "issuer": "Civitas 政策仿真系统",
        "classification": "内部研判材料",
        "recipient": recipient,
        "title": title,
    }


def _parse_markdown(markdown_text: str) -> List[Tuple[str, str]]:
    blocks: List[Tuple[str, str]] = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            blocks.append(("blank", ""))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        else:
            blocks.append(("text", line.strip()))
    return blocks


def _content_blocks(markdown_text: str, report_meta: Dict[str, str]) -> List[Tuple[str, str]]:
    blocks = _parse_markdown(markdown_text)
    if blocks and blocks[0][0] == "h1" and blocks[0][1] == report_meta["title"]:
        return blocks[1:]
    return blocks


def _set_run_font(run, size_pt: float, *, bold: bool = False, color: str | None = None) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _append_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _build_docx_cover(document: Document, report_meta: Dict[str, str]) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title_run = title.add_run(report_meta["title"])
    _set_run_font(title_run, 22, bold=True, color="17304F")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("正式汇报材料")
    _set_run_font(subtitle_run, 14, bold=True, color="305178")

    meta_table = document.add_table(rows=5, cols=2)
    meta_table.style = "Table Grid"
    rows = [
        ("报告编号", report_meta["report_no"]),
        ("报送对象", report_meta["recipient"]),
        ("材料性质", report_meta["classification"]),
        ("生成日期", report_meta["date_cn"]),
        ("出具单位", report_meta["issuer"]),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = meta_table.rows[idx].cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, 11)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    note_run = note.add_run("本材料用于政策研判、会商沟通和正式汇报底稿。")
    _set_run_font(note_run, 11, color="4B6587")


def _build_docx_body(document: Document, markdown_text: str, report_meta: Dict[str, str]) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(24)
    section.right_margin = Mm(20)
    section.different_first_page_header_footer = True

    _build_docx_cover(document, report_meta)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    content_section.page_width = Mm(210)
    content_section.page_height = Mm(297)
    content_section.top_margin = Mm(22)
    content_section.bottom_margin = Mm(18)
    content_section.left_margin = Mm(24)
    content_section.right_margin = Mm(20)

    header = content_section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(f"{report_meta['issuer']}    {report_meta['report_no']}")
    _set_run_font(header_run, 9, color="4B6587")

    footer = content_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"{report_meta['classification']}  |  第 ")
    _set_run_font(footer_run, 9, color="4B6587")
    _append_page_number(footer)
    footer_tail = footer.add_run(" 页")
    _set_run_font(footer_tail, 9, color="4B6587")

    for kind, text in _content_blocks(markdown_text, report_meta):
        if kind == "blank":
            document.add_paragraph()
            continue

        paragraph = document.add_paragraph()
        if kind == "h1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            _set_run_font(run, 18, bold=True, color="17304F")
        elif kind == "h2":
            paragraph.paragraph_format.space_before = Pt(10)
            run = paragraph.add_run(text)
            _set_run_font(run, 14, bold=True, color="17304F")
        elif kind == "h3":
            run = paragraph.add_run(text)
            _set_run_font(run, 12, bold=True, color="305178")
        elif kind == "bullet":
            paragraph.style = document.styles["List Bullet"]
            run = paragraph.add_run(text)
            _set_run_font(run, 10.5)
        else:
            paragraph.paragraph_format.first_line_indent = Pt(21)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(text)
            _set_run_font(run, 10.5)


def _build_docx_bytes(markdown_text: str, report_meta: Dict[str, str]) -> bytes:
    document = Document()
    _build_docx_body(document, markdown_text, report_meta)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_styles():
    registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    return {
        "cover_title": ParagraphStyle(
            "CoverTitle",
            parent=styles["Title"],
            fontName="STSong-Light",
            fontSize=22,
            leading=28,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#17304F"),
        ),
        "cover_subtitle": ParagraphStyle(
            "CoverSubtitle",
            parent=styles["Normal"],
            fontName="STSong-Light",
            fontSize=13,
            leading=18,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#305178"),
        ),
        "h1": ParagraphStyle(
            "ReportH1",
            parent=styles["Heading1"],
            fontName="STSong-Light",
            fontSize=17,
            leading=24,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#17304F"),
        ),
        "h2": ParagraphStyle(
            "ReportH2",
            parent=styles["Heading2"],
            fontName="STSong-Light",
            fontSize=13,
            leading=19,
            spaceBefore=10,
            spaceAfter=4,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#17304F"),
        ),
        "h3": ParagraphStyle(
            "ReportH3",
            parent=styles["Heading3"],
            fontName="STSong-Light",
            fontSize=11.5,
            leading=16,
            spaceBefore=6,
            spaceAfter=2,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#305178"),
        ),
        "body": ParagraphStyle(
            "ReportBody",
            parent=styles["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=16,
            alignment=TA_JUSTIFY,
            firstLineIndent=21,
            textColor=colors.HexColor("#1F2B44"),
        ),
        "bullet": ParagraphStyle(
            "ReportBullet",
            parent=styles["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=16,
            leftIndent=14,
            firstLineIndent=-10,
            bulletIndent=0,
            textColor=colors.HexColor("#1F2B44"),
        ),
        "meta": ParagraphStyle(
            "ReportMeta",
            parent=styles["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=15,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#1F2B44"),
        ),
    }


def _draw_pdf_header_footer(canvas, doc, report_meta: Dict[str, str]) -> None:
    if doc.page == 1:
        return
    canvas.saveState()
    canvas.setFont("STSong-Light", 9)
    canvas.setFillColor(colors.HexColor("#4B6587"))
    canvas.drawString(doc.leftMargin, A4[1] - 14 * mm, report_meta["issuer"])
    canvas.drawRightString(A4[0] - doc.rightMargin, A4[1] - 14 * mm, report_meta["report_no"])
    canvas.drawString(doc.leftMargin, 10 * mm, report_meta["classification"])
    canvas.drawCentredString(A4[0] / 2, 10 * mm, f"第 {canvas.getPageNumber() - 1} 页")
    canvas.restoreState()


def _build_pdf_bytes(markdown_text: str, report_meta: Dict[str, str]) -> bytes:
    styles = _pdf_styles()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=24 * mm,
        rightMargin=20 * mm,
        topMargin=22 * mm,
        bottomMargin=18 * mm,
    )

    story: List[Any] = []
    story.append(Spacer(1, 50 * mm))
    story.append(Paragraph(report_meta["title"], styles["cover_title"]))
    story.append(Spacer(1, 8 * mm))
    story.append(Paragraph("正式汇报材料", styles["cover_subtitle"]))
    story.append(Spacer(1, 18 * mm))

    cover_table = Table(
        [
            ["报告编号", report_meta["report_no"]],
            ["报送对象", report_meta["recipient"]],
            ["材料性质", report_meta["classification"]],
            ["生成日期", report_meta["date_cn"]],
            ["出具单位", report_meta["issuer"]],
        ],
        colWidths=[32 * mm, 118 * mm],
    )
    cover_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 10.5),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1F2B44")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EDF2F9")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C9D8EF")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(cover_table)
    story.append(Spacer(1, 16 * mm))
    story.append(Paragraph("本材料用于政策研判、会商沟通和正式汇报底稿。", styles["cover_subtitle"]))
    story.append(PageBreak())

    for kind, text in _content_blocks(markdown_text, report_meta):
        if kind == "blank":
            story.append(Spacer(1, 3 * mm))
        elif kind == "h1":
            story.append(Paragraph(text, styles["h1"]))
        elif kind == "h2":
            story.append(Paragraph(text, styles["h2"]))
        elif kind == "h3":
            story.append(Paragraph(text, styles["h3"]))
        elif kind == "bullet":
            story.append(Paragraph(f"• {text}", styles["bullet"]))
        else:
            story.append(Paragraph(text, styles["body"]))

    doc.build(
        story,
        onFirstPage=lambda canvas, doc: None,
        onLaterPages=lambda canvas, doc: _draw_pdf_header_footer(canvas, doc, report_meta),
    )
    return buffer.getvalue()


def write_report_artifacts(
    *,
    root_dir: Path,
    report_type: str,
    title: str,
    markdown_text: str,
    payload: Dict[str, Any],
) -> Dict[str, Any]:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    root_dir.mkdir(parents=True, exist_ok=True)
    stem = f"{report_type}_{safe_slug(title, fallback=report_type)}_{timestamp}"
    markdown_path = root_dir / f"{stem}.md"
    json_path = root_dir / f"{stem}.json"
    docx_path = root_dir / f"{stem}.docx"
    pdf_path = root_dir / f"{stem}.pdf"

    clean_payload = serializable_payload(payload)
    report_meta = clean_payload.get("report_meta") or official_report_meta(report_type, title)

    markdown_path.write_text(markdown_text, encoding="utf-8")
    json_text = json.dumps(clean_payload, ensure_ascii=False, indent=2)
    json_path.write_text(json_text, encoding="utf-8")

    docx_bytes = _build_docx_bytes(markdown_text, report_meta)
    docx_path.write_bytes(docx_bytes)

    pdf_bytes = _build_pdf_bytes(markdown_text, report_meta)
    pdf_path.write_bytes(pdf_bytes)

    return {
        "stem": stem,
        "timestamp": timestamp,
        "report_meta": report_meta,
        "markdown_path": markdown_path,
        "json_path": json_path,
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "markdown_text": markdown_text,
        "json_text": json_text,
        "docx_bytes": docx_bytes,
        "pdf_bytes": pdf_bytes,
    }
